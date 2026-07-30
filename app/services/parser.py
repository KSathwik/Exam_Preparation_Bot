"""Document parsing module for PDF and DOCX files."""

import re
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

import pdfplumber
from docx import Document as DocxDocument
from loguru import logger

from app.core.config import settings

from .models import ChunkMetadata, Document, DocumentChunk


@dataclass
class _Block:
    """A structurally-scoped span of text (one heading section, or one page
    when no reliable heading signal exists) — chunked independently so a
    single chunk never straddles a section/heading boundary."""

    text: str
    section_title: Optional[str] = None
    section_level: Optional[int] = None
    page_number: Optional[int] = None


class DocumentParser:
    """Parse documents and extract text with metadata."""

    def __init__(self, max_chunk_size: int = None, chunk_overlap: int = None):
        self.max_chunk_size = max_chunk_size or settings.max_chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
        self.min_chunk_size = settings.min_chunk_size

    def parse_document(self, file_path: str, file_type: str) -> Document:
        if file_type.lower() == "pdf":
            return self._parse_pdf(file_path)
        elif file_type.lower() == "docx":
            return self._parse_docx(file_path)
        raise ValueError(f"Unsupported file type: {file_type}")

    def _parse_pdf(self, file_path: str) -> Document:
        logger.info(f"[PARSER] Parsing PDF: {file_path}")
        file_size_bytes = Path(file_path).stat().st_size
        file_name = Path(file_path).name

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"[PARSER] PDF opened: pages={total_pages}  size={file_size_bytes/1024:.1f} KB")
            blocks: List[_Block] = []
            for page_num, page in enumerate(pdf.pages, 1):
                page_blocks = _pdf_page_to_blocks(page, page_num)
                logger.debug(f"[PARSER] Page {page_num}: {len(page_blocks)} block(s) extracted")
                blocks.extend(page_blocks)

            total_chars = sum(len(b.text) for b in blocks)
            logger.info(f"[PARSER] Total text extracted: {total_chars} chars across {len(blocks)} block(s)")
            chunks = self._create_chunks(blocks, file_name, total_pages)

        logger.info(f"[PARSER] PDF parsed: {file_name}  chunks={len(chunks)}")
        return Document(
            file_name=file_name,
            file_type="pdf",
            file_size_bytes=file_size_bytes,
            total_chunks=len(chunks),
            chunks=chunks,
            upload_timestamp=datetime.now().isoformat(),
        )

    def _parse_docx(self, file_path: str) -> Document:
        logger.info(f"[PARSER] Parsing DOCX: {file_path}")
        file_size_bytes = Path(file_path).stat().st_size
        file_name = Path(file_path).name

        doc = DocxDocument(file_path)
        blocks = _docx_to_blocks(doc)
        total_chars = sum(len(b.text) for b in blocks)
        logger.info(
            f"[PARSER] DOCX text extracted: {total_chars} chars  "
            f"paragraphs={len(doc.paragraphs)}  blocks={len(blocks)}"
        )
        chunks = self._create_chunks(blocks, file_name, len(doc.paragraphs))

        logger.info(f"[PARSER] DOCX parsed: {file_name}  chunks={len(chunks)}")
        return Document(
            file_name=file_name,
            file_type="docx",
            file_size_bytes=file_size_bytes,
            total_chunks=len(chunks),
            chunks=chunks,
            upload_timestamp=datetime.now().isoformat(),
        )

    def _create_chunks(self, blocks: List[_Block], file_name: str, total_pages: int) -> List[DocumentChunk]:
        """Accumulate each block's sentences into word-count-bounded chunks.

        Chunking resets at every block boundary (never straddles a
        heading/page split), but a trailing fragment under ``min_chunk_size``
        is never dropped: it merges into the previous chunk from the same
        block if one exists, otherwise carries forward into the next block's
        accumulation, otherwise (no next block either) is force-emitted as
        its own chunk. A non-empty document therefore never yields zero
        chunks — the fragment always lands somewhere.

        Chunk content is collected as plain (text, source_block) pairs first
        and only turned into ``DocumentChunk``s once the final count is
        known — ``chunk_position`` and the page-ratio fallback both depend on
        ``total_chunks``, so computing them against a placeholder count (as
        the previous implementation did) silently produced wrong metadata
        for every chunk after the first.
        """
        pending: List[Tuple[str, _Block]] = []
        carry_sentences: List[str] = []

        for block in blocks:
            sentences = carry_sentences + _split_sentences(block.text)
            carry_sentences = []
            if not sentences:
                continue

            block_start = len(pending)
            current_chunk = ""
            chunk_sentences: List[str] = []

            for sentence in sentences:
                test_chunk = f"{current_chunk} {sentence}" if current_chunk else sentence
                if len(test_chunk.split()) > self.max_chunk_size:
                    if current_chunk and len(current_chunk.split()) >= self.min_chunk_size:
                        pending.append((current_chunk.strip(), block))
                        overlap = chunk_sentences[-2:] if len(chunk_sentences) > 1 else []
                        current_chunk = " ".join(overlap)
                        chunk_sentences = list(overlap)
                current_chunk = f"{current_chunk} {sentence}" if current_chunk else sentence
                chunk_sentences.append(sentence)

            current_chunk = current_chunk.strip()
            if not current_chunk:
                continue
            if len(current_chunk.split()) >= self.min_chunk_size:
                pending.append((current_chunk, block))
            elif len(pending) > block_start:
                # Merge the undersized trailing fragment into the last chunk
                # emitted for *this* block rather than dropping it.
                last_content, last_block = pending[-1]
                pending[-1] = (f"{last_content} {current_chunk}".strip(), last_block)
            else:
                # This block produced no chunk of its own yet — carry the
                # fragment forward so the next block's accumulation absorbs
                # it instead of losing it.
                carry_sentences = _split_sentences(current_chunk)

        if carry_sentences:
            leftover = " ".join(carry_sentences).strip()
            if leftover:
                if pending:
                    last_content, last_block = pending[-1]
                    pending[-1] = (f"{last_content} {leftover}".strip(), last_block)
                else:
                    # Nothing was ever emitted (the whole document is smaller
                    # than min_chunk_size) — force-emit it as the sole chunk
                    # so a non-empty document never yields zero chunks.
                    pending.append((leftover, blocks[-1] if blocks else _Block(text=leftover)))

        total = len(pending)
        return [
            _make_chunk(
                content,
                file_name,
                idx,
                total,
                total_pages,
                section_title=block.section_title,
                section_level=block.section_level,
                page_number_override=block.page_number,
            )
            for idx, (content, block) in enumerate(pending)
        ]


# ----------------------------------------------------------------------
# DOCX structural blocks
# ----------------------------------------------------------------------

_DOCX_HEADING_RE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)


def _docx_heading_level(style_name: Optional[str]) -> Optional[int]:
    if not style_name:
        return None
    name = style_name.strip()
    match = _DOCX_HEADING_RE.match(name)
    if match:
        return int(match.group(1))
    if name.lower() == "title":
        return 1
    return None


def _docx_to_blocks(doc: DocxDocument) -> List[_Block]:
    """Split DOCX paragraphs into heading-scoped blocks using
    ``paragraph.style.name`` — an exact signal, no heuristic needed."""
    blocks: List[_Block] = []
    current_title: Optional[str] = None
    current_level: Optional[int] = None
    current_lines: List[str] = []

    def flush() -> None:
        text = "\n".join(current_lines).strip()
        if text:
            blocks.append(_Block(text=text, section_title=current_title, section_level=current_level))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        level = _docx_heading_level(para.style.name if para.style is not None else None)
        if level is not None:
            flush()
            current_lines = []
            current_title = text
            current_level = level
            continue
        current_lines.append(text)

    flush()
    return blocks


# ----------------------------------------------------------------------
# PDF structural blocks (font-size heading heuristic)
# ----------------------------------------------------------------------

_HEADING_SIZE_RATIO = 1.15
_HEADING_MAX_WORDS = 12


def _pdf_page_to_blocks(page, page_number: int) -> List[_Block]:
    """Split one PDF page into heading-scoped blocks via a font-size
    heuristic, falling back to one block per page when there's no reliable
    signal (no extractable words, or every line is the same size)."""
    try:
        words = page.extract_words(extra_attrs=["size", "fontname"])
    except Exception as exc:
        logger.debug(f"[PARSER] extract_words failed on page {page_number}: {exc}")
        words = []

    sizes = [w["size"] for w in words if w.get("size")]
    if not words or not sizes:
        text = (page.extract_text() or "").strip()
        return [_Block(text=text, page_number=page_number)] if text else []

    lines: List[List[dict]] = []
    for word in sorted(words, key=lambda w: (round(w["top"], 1), w["x0"])):
        if lines and abs(lines[-1][0]["top"] - word["top"]) < 2:
            lines[-1].append(word)
        else:
            lines.append([word])

    body_size = Counter(round(s) for s in sizes).most_common(1)[0][0]
    heading_threshold = body_size * _HEADING_SIZE_RATIO

    blocks: List[_Block] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []

    def flush() -> None:
        text = " ".join(current_lines).strip()
        if text:
            blocks.append(_Block(text=text, section_title=current_title, page_number=page_number))

    for line_words in lines:
        line_text = " ".join(w["text"] for w in line_words).strip()
        if not line_text:
            continue
        avg_size = sum(w["size"] for w in line_words) / len(line_words)
        # Below this confidence bar we don't guess a heading — a wrong
        # section_title actively misleads citations, worse than "unknown".
        is_heading = avg_size >= heading_threshold and len(line_text.split()) <= _HEADING_MAX_WORDS
        if is_heading:
            flush()
            current_lines = []
            current_title = line_text
            continue
        current_lines.append(line_text)

    flush()
    return blocks


def _make_chunk(
    content: str,
    file_name: str,
    chunk_index: int,
    total_chunks: int,
    total_pages: int,
    section_title: Optional[str] = None,
    section_level: Optional[int] = None,
    page_number_override: Optional[int] = None,
) -> DocumentChunk:
    safe_total = max(total_chunks, 1)
    page_number = page_number_override or (max(1, int((chunk_index / safe_total) * total_pages) + 1))
    if safe_total <= 3:
        position = "unknown"
    elif chunk_index < safe_total * 0.33:
        position = "beginning"
    elif chunk_index > safe_total * 0.66:
        position = "end"
    else:
        position = "middle"

    return DocumentChunk(
        content=content,
        metadata=ChunkMetadata(
            page_number=min(page_number, total_pages) if total_pages else 1,
            section_title=section_title,
            section_level=section_level,
            chunk_index=chunk_index,
            total_chunks=total_chunks,
            file_name=file_name,
            chunk_position=position,
        ),
    )


def _split_sentences(text: str) -> List[str]:
    parts = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in parts if s.strip()]


def parse_file(file_path: str) -> Document:
    """Convenience function to parse a file."""
    ext = Path(file_path).suffix.lower().lstrip(".")
    if ext not in ("pdf", "docx"):
        raise ValueError(f"Unsupported file type: {ext}")
    return DocumentParser().parse_document(file_path, ext)
