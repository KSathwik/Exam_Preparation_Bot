"""Tests for document parser utilities (no real files needed)."""

from unittest.mock import MagicMock

import pytest
from docx import Document as DocxDocument

from app.core.config import settings
from app.services.parser import (
    DocumentParser,
    _Block,
    _docx_to_blocks,
    _make_chunk,
    _pdf_page_to_blocks,
    _split_sentences,
)


def test_split_sentences_basic():
    text = "First sentence. Second sentence! Third sentence?"
    result = _split_sentences(text)
    assert len(result) == 3
    assert result[0] == "First sentence."


def test_split_sentences_empty():
    assert _split_sentences("") == []


def test_split_sentences_no_punctuation():
    result = _split_sentences("Just some text without a period")
    assert len(result) == 1


def test_make_chunk_beginning_position():
    chunk = _make_chunk("Some content", "test.pdf", chunk_index=0, total_chunks=10, total_pages=5)
    assert chunk.metadata.chunk_position == "beginning"
    assert chunk.metadata.file_name == "test.pdf"


def test_make_chunk_end_position():
    chunk = _make_chunk("Some content", "test.pdf", chunk_index=9, total_chunks=10, total_pages=5)
    assert chunk.metadata.chunk_position == "end"


def test_make_chunk_middle_position():
    chunk = _make_chunk("Some content", "test.pdf", chunk_index=5, total_chunks=10, total_pages=5)
    assert chunk.metadata.chunk_position == "middle"


def test_make_chunk_unknown_when_few():
    chunk = _make_chunk("Content", "test.pdf", chunk_index=1, total_chunks=2, total_pages=2)
    assert chunk.metadata.chunk_position == "unknown"


def test_make_chunk_page_number_bounded():
    chunk = _make_chunk("Content", "test.pdf", chunk_index=99, total_chunks=100, total_pages=5)
    assert chunk.metadata.page_number <= 5


def test_make_chunk_page_number_override_takes_priority():
    chunk = _make_chunk(
        "Content", "test.pdf", chunk_index=0, total_chunks=1, total_pages=20, page_number_override=17
    )
    assert chunk.metadata.page_number == 17


def test_make_chunk_carries_section_title_and_level():
    chunk = _make_chunk(
        "Content",
        "test.pdf",
        chunk_index=0,
        total_chunks=1,
        total_pages=1,
        section_title="Chapter One",
        section_level=1,
    )
    assert chunk.metadata.section_title == "Chapter One"
    assert chunk.metadata.section_level == 1


# ── DOCX structural blocks (heading detection via paragraph.style.name) ──


def test_docx_to_blocks_splits_on_headings():
    doc = DocxDocument()
    doc.add_paragraph("Intro text before any heading.")
    doc.add_heading("Section One", level=1)
    doc.add_paragraph("Body of section one.")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Body of section two.")

    blocks = _docx_to_blocks(doc)

    assert len(blocks) == 3
    assert blocks[0].section_title is None
    assert blocks[0].text == "Intro text before any heading."
    assert blocks[1].section_title == "Section One"
    assert blocks[1].section_level == 1
    assert blocks[1].text == "Body of section one."
    assert blocks[2].section_title == "Section Two"
    assert blocks[2].section_level == 2
    assert blocks[2].text == "Body of section two."


def test_docx_to_blocks_no_headings_yields_single_block():
    doc = DocxDocument()
    doc.add_paragraph("Just one plain paragraph.")
    doc.add_paragraph("And another plain paragraph.")

    blocks = _docx_to_blocks(doc)

    assert len(blocks) == 1
    assert blocks[0].section_title is None
    assert "Just one plain paragraph." in blocks[0].text
    assert "And another plain paragraph." in blocks[0].text


def test_docx_to_blocks_empty_document_yields_no_blocks():
    doc = DocxDocument()
    assert _docx_to_blocks(doc) == []


# ── PDF structural blocks (font-size heading heuristic) ──────────────


def _word(text, top, x0, size):
    return {"text": text, "top": top, "x0": x0, "size": size, "fontname": "Test"}


def test_pdf_page_to_blocks_detects_heading_by_font_size():
    page = MagicMock()
    page.extract_words.return_value = [
        _word("Chapter", 0, 0, 20),
        _word("One", 0, 50, 20),
        _word("Body", 20, 0, 10),
        _word("text", 20, 40, 10),
        _word("here.", 20, 80, 10),
    ]

    blocks = _pdf_page_to_blocks(page, page_number=3)

    assert len(blocks) == 1
    assert blocks[0].section_title == "Chapter One"
    assert blocks[0].page_number == 3
    assert blocks[0].text == "Body text here."


def test_pdf_page_to_blocks_falls_back_to_full_page_when_no_words():
    page = MagicMock()
    page.extract_words.return_value = []
    page.extract_text.return_value = "Plain text with no font info."

    blocks = _pdf_page_to_blocks(page, page_number=1)

    assert len(blocks) == 1
    assert blocks[0].section_title is None
    assert blocks[0].text == "Plain text with no font info."


def test_pdf_page_to_blocks_falls_back_when_extract_words_raises():
    page = MagicMock()
    page.extract_words.side_effect = Exception("boom")
    page.extract_text.return_value = "Recovered text."

    blocks = _pdf_page_to_blocks(page, page_number=2)

    assert len(blocks) == 1
    assert blocks[0].text == "Recovered text."


def test_pdf_page_to_blocks_empty_page_returns_no_blocks():
    page = MagicMock()
    page.extract_words.return_value = []
    page.extract_text.return_value = ""

    assert _pdf_page_to_blocks(page, page_number=1) == []


# ── DocumentParser._create_chunks: structural awareness + no-data-loss ──


def test_create_chunks_force_emits_undersized_whole_document():
    """A non-empty document must never yield zero chunks, even if its
    entire content is under min_chunk_size."""
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 50
    try:
        parser = DocumentParser(max_chunk_size=1000, chunk_overlap=0)
        chunks = parser._create_chunks([_Block(text="Short sentence here.")], "doc.pdf", total_pages=1)
    finally:
        settings.min_chunk_size = original_min

    assert len(chunks) == 1
    assert chunks[0].content == "Short sentence here."
    assert chunks[0].metadata.total_chunks == 1


def test_create_chunks_merges_trailing_fragment_into_previous_chunk():
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 5
    try:
        parser = DocumentParser(max_chunk_size=6, chunk_overlap=0)
        text = "One two three four five six seven. Eight nine."
        chunks = parser._create_chunks([_Block(text=text)], "doc.pdf", total_pages=1)
    finally:
        settings.min_chunk_size = original_min

    assert len(chunks) == 1
    assert "Eight nine" in chunks[0].content


def test_create_chunks_carries_trailing_fragment_into_next_block():
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 5
    try:
        parser = DocumentParser(max_chunk_size=1000, chunk_overlap=0)
        blocks = [
            _Block(text="Tiny bit."),
            _Block(text="This block has enough words to meet the minimum easily today."),
        ]
        chunks = parser._create_chunks(blocks, "doc.pdf", total_pages=1)
    finally:
        settings.min_chunk_size = original_min

    assert len(chunks) == 1
    assert chunks[0].content.startswith("Tiny bit.")


def test_create_chunks_never_returns_empty_for_nonempty_blocks():
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 100
    try:
        parser = DocumentParser(max_chunk_size=1000, chunk_overlap=0)
        blocks = [_Block(text="a."), _Block(text="b."), _Block(text="c.")]
        chunks = parser._create_chunks(blocks, "doc.pdf", total_pages=1)
    finally:
        settings.min_chunk_size = original_min

    assert len(chunks) == 1
    assert chunks[0].content == "a. b. c."


def test_create_chunks_resets_at_block_boundaries():
    """Each block is accumulated independently — a single chunk never
    straddles a heading/page split when a block alone clears min_chunk_size."""
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 2
    try:
        parser = DocumentParser(max_chunk_size=1000, chunk_overlap=0)
        blocks = [
            _Block(text="First section body text here.", section_title="Intro"),
            _Block(text="Second section body text here.", section_title="Details"),
        ]
        chunks = parser._create_chunks(blocks, "doc.pdf", total_pages=1)
    finally:
        settings.min_chunk_size = original_min

    assert len(chunks) == 2
    assert chunks[0].metadata.section_title == "Intro"
    assert chunks[1].metadata.section_title == "Details"


def test_create_chunks_assigns_varied_positions_not_always_unknown():
    """Regression test: chunk_position/total_chunks used to be computed
    against a placeholder total_chunks=0 at creation time and never
    recomputed, so every real chunk silently got chunk_position='unknown'."""
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 2
    try:
        parser = DocumentParser(max_chunk_size=3, chunk_overlap=0)
        text = " ".join(f"word{i}." for i in range(40))
        chunks = parser._create_chunks([_Block(text=text)], "doc.docx", total_pages=10)
    finally:
        settings.min_chunk_size = original_min

    assert len(chunks) > 3
    positions = {c.metadata.chunk_position for c in chunks}
    assert positions != {"unknown"}
    assert all(c.metadata.total_chunks == len(chunks) for c in chunks)


def test_create_chunks_docx_page_numbers_increase_with_chunk_index():
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 2
    try:
        parser = DocumentParser(max_chunk_size=3, chunk_overlap=0)
        text = " ".join(f"word{i}." for i in range(40))
        chunks = parser._create_chunks([_Block(text=text)], "doc.docx", total_pages=10)
    finally:
        settings.min_chunk_size = original_min

    pages = [c.metadata.page_number for c in chunks]
    assert pages == sorted(pages)
    assert pages[0] == 1
    assert pages[-1] <= 10
    assert len(set(pages)) > 1


def test_create_chunks_pdf_blocks_use_real_page_number_override():
    original_min = settings.min_chunk_size
    settings.min_chunk_size = 2
    try:
        parser = DocumentParser(max_chunk_size=1000, chunk_overlap=0)
        blocks = [
            _Block(text="Text on page five.", page_number=5),
            _Block(text="Text on page twelve.", page_number=12),
        ]
        chunks = parser._create_chunks(blocks, "doc.pdf", total_pages=20)
    finally:
        settings.min_chunk_size = original_min

    assert len(chunks) == 2
    assert chunks[0].metadata.page_number == 5
    assert chunks[1].metadata.page_number == 12
